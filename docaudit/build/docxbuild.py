#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Builder for the Vietnamese doctoral monograph (chuyên đề).

Line-oriented source markup -> .docx with:
  - Times New Roman for ascii / hAnsi / eastAsia / cs on every run and style
  - four heading levels (CHƯƠNG / x.y / x.y.z / x.y.z.w)
  - identical caption style for tables (above) and figures (below), with the
    citation bracket sitting immediately after the caption text (never "Nguồn:")
  - balanced tables: fixed layout, widths summing to the 15.5 cm text width,
    repeating shaded header row, rows that cannot split across pages
  - front matter in roman numerals, body restarting at 1
  - auto TOC field plus lists of tables/figures built from PAGEREF bookmarks
"""

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"
BODY_PT = 13
TABLE_PT = 11.5
TEXT_WIDTH_CM = 15.5  # 21.0 - 3.5 - 2.0
BLACK = RGBColor(0, 0, 0)


# ---------------------------------------------------------------- run helpers
def _rfonts(rpr):
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)
    return rf


def set_run_font(run, size=BODY_PT, bold=None, italic=None):
    f = run.font
    f.name = FONT
    f.size = Pt(size)
    f.color.rgb = BLACK
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    rpr = run._element.get_or_add_rPr()
    _rfonts(rpr)
    # complex-script size, inserted immediately after w:sz to keep child order
    sz = rpr.find(qn("w:sz"))
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        if sz is not None:
            sz.addnext(szcs)
        else:
            rpr.append(szcs)
    szcs.set(qn("w:val"), str(int(round(size * 2))))
    return run


def add_run(par, text, size=BODY_PT, bold=None, italic=None):
    return set_run_font(par.add_run(text), size, bold, italic)


_TOKEN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\^\{[^}]*\}|_\{[^}]*\})")


def add_rich(par, text, size=BODY_PT, bold=None, italic=None):
    """Inline markers: **bold**, *italic*, ^{superscript}, _{subscript}."""
    pos = 0
    for m in _TOKEN.finditer(text):
        if m.start() > pos:
            add_run(par, text[pos:m.start()], size, bold, italic)
        tok = m.group(0)
        if tok.startswith("**"):
            add_run(par, tok[2:-2], size, True, italic)
        elif tok.startswith("*"):
            add_run(par, tok[1:-1], size, bold, True)
        elif tok.startswith("^"):
            set_run_font(par.add_run(tok[2:-1]), size, bold, italic).font.superscript = True
        else:
            set_run_font(par.add_run(tok[2:-1]), size, bold, italic).font.subscript = True
        pos = m.end()
    if pos < len(text):
        add_run(par, text[pos:], size, bold, italic)


def field(par, instr, cached, size=BODY_PT, bold=None, italic=None):
    """Complex field with a cached result so it reads sensibly before updating."""
    def _fc(kind):
        r = par.add_run()
        set_run_font(r, size, bold, italic)
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        r._element.append(el)

    _fc("begin")
    r = par.add_run()
    set_run_font(r, size, bold, italic)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r._element.append(it)
    _fc("separate")
    add_run(par, cached, size, bold, italic)
    _fc("end")


def bookmark(par, name, bid):
    s = OxmlElement("w:bookmarkStart")
    s.set(qn("w:id"), str(bid))
    s.set(qn("w:name"), name)
    e = OxmlElement("w:bookmarkEnd")
    e.set(qn("w:id"), str(bid))
    par._p.insert(0, s)
    par._p.append(e)


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def cell_margins(table, top=40, bottom=40, left=85, right=85):
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement("w:" + name)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    table._tbl.tblPr.append(mar)


def fixed_layout(table):
    el = OxmlElement("w:tblLayout")
    el.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(el)


def row_flag(row, tag):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:" + tag)
    trpr.append(el)


def set_col_widths(table, widths_cm):
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths_cm):
            gc.set(qn("w:w"), str(int(round(w * 567))))
    for row in table.rows:
        for c, w in zip(row.cells, widths_cm):
            c.width = Cm(w)


# ------------------------------------------------------------------ document
def new_document():
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BODY_PT)
    st.font.color.rgb = BLACK
    _rfonts(st.element.get_or_add_rPr())
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for s in doc.styles:
        try:
            if s.type is not None and s.element.tag.endswith("}style"):
                rpr = s.element.find(qn("w:rPr"))
                if rpr is None:
                    rpr = s.element.get_or_add_rPr()
                _rfonts(rpr)
        except Exception:
            pass

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(3.0), Cm(3.0)
    sec.left_margin, sec.right_margin = Cm(3.5), Cm(2.0)
    sec.header_distance, sec.footer_distance = Cm(1.5), Cm(1.5)

    h1 = doc.styles["Heading 1"]
    h1.font.size, h1.font.bold, h1.font.italic = Pt(14), True, False
    h1.font.color.rgb = BLACK
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(14)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.keep_with_next = True

    for name, size, bold, italic, sb, sa in (
        ("Heading 2", 13, True, False, 12, 6),
        ("Heading 3", 13, True, False, 10, 6),
        ("Heading 4", 13, True, True, 8, 6),
    ):
        s = doc.styles[name]
        s.font.size, s.font.bold, s.font.italic = Pt(size), bold, italic
        s.font.color.rgb = BLACK
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after = Pt(sa)
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.keep_with_next = True

    settings = doc.settings.element
    uf = settings.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings.append(uf)
    uf.set(qn("w:val"), "true")
    return doc


def section_page_numbers(section, fmt, start=None, show=True):
    """fmt: 'lowerRoman' | 'decimal'."""
    sectPr = section._sectPr
    old = sectPr.find(qn("w:pgNumType"))
    if old is not None:
        sectPr.remove(old)
    el = OxmlElement("w:pgNumType")
    el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    sectPr.append(el)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if show:
        field(p, " PAGE ", "1", size=BODY_PT)


# ------------------------------------------------------------------- builder
class Builder:
    def __init__(self, doc, media_dir):
        self.doc = doc
        self.media = media_dir
        self.chapter = 0
        self.table_no = 0
        self.figure_no = 0
        self.tables = []
        self.figures = []
        self._bid = 1000

    def _next_bid(self):
        self._bid += 1
        return self._bid

    # -- headings ----------------------------------------------------------
    def heading(self, level, text, page_break=False):
        if page_break:
            self.doc.add_page_break()
        p = self.doc.add_paragraph(style="Heading %d" % level)
        p.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        )
        p.paragraph_format.first_line_indent = Cm(0)
        add_rich(p, text, size=(14 if level == 1 else 13), bold=True, italic=(level == 4))
        if level == 1:
            m = re.match(r"CHƯƠNG\s+(\d+)", text)
            if m:
                self.chapter = int(m.group(1))
                self.table_no = 0
                self.figure_no = 0
        return p

    # -- body --------------------------------------------------------------
    def para(self, text, indent=True, align="j", size=BODY_PT, space_after=6, bold=False, italic=False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = {
            "j": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "c": WD_ALIGN_PARAGRAPH.CENTER,
            "l": WD_ALIGN_PARAGRAPH.LEFT,
            "r": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        pf.first_line_indent = Cm(1.0) if indent else Cm(0)
        pf.line_spacing = 1.5
        pf.space_after = Pt(space_after)
        add_rich(p, text, size=size, bold=bold or None, italic=italic or None)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-0.5)
        pf.line_spacing = 1.5
        pf.space_after = Pt(4)
        add_rich(p, "– " + text)
        return p

    def reference(self, text):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        pf.line_spacing = 1.2
        pf.space_after = Pt(6)
        add_rich(p, text, size=BODY_PT)
        return p

    # -- captions ----------------------------------------------------------
    def _caption(self, kind, text, above):
        if kind == "Bảng":
            self.table_no += 1
            n = self.table_no
        else:
            self.figure_no += 1
            n = self.figure_no
        label = "%s %d.%d." % (kind, self.chapter, n)
        bname = "_%s_%d_%d" % ("Bang" if kind == "Bảng" else "Hinh", self.chapter, n)

        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = 1.25
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0.5)
        pf.right_indent = Cm(0.5)
        pf.space_before = Pt(6 if above else 4)
        pf.space_after = Pt(4 if above else 10)
        pf.keep_with_next = above
        add_run(p, label + " ", size=BODY_PT, bold=True, italic=False)
        add_rich(p, text, size=BODY_PT, bold=False, italic=False)
        bookmark(p, bname, self._next_bid())

        (self.tables if kind == "Bảng" else self.figures).append((label, text, bname))
        return p

    # -- tables ------------------------------------------------------------
    def table(self, caption, header, rows, widths=None, aligns=None, fontsize=None):
        if caption:
            self._caption("Bảng", caption, above=True)
        ncol = len(header)
        t = self.doc.add_table(rows=1, cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        fixed_layout(t)
        cell_margins(t)

        if not widths:
            widths = [1.0] * ncol
        widths = (list(widths) + [1.0] * ncol)[:ncol]
        tot = float(sum(widths))
        wcm = [TEXT_WIDTH_CM * w / tot for w in widths]
        fs = fontsize or TABLE_PT

        def fill(cells, values, is_header):
            for i, (c, v) in enumerate(zip(cells, values)):
                p = c.paragraphs[0]
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing = 1.05
                pf.first_line_indent = Cm(0)
                if is_header:
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    a = aligns[i] if aligns and i < len(aligns) else ("l" if i == 0 else "c")
                    pf.alignment = {
                        "l": WD_ALIGN_PARAGRAPH.LEFT,
                        "c": WD_ALIGN_PARAGRAPH.CENTER,
                        "j": WD_ALIGN_PARAGRAPH.JUSTIFY,
                        "r": WD_ALIGN_PARAGRAPH.RIGHT,
                    }[a]
                add_rich(p, v, size=fs, bold=(True if is_header else None))
                c.vertical_alignment = 1

        hdr = t.rows[0]
        fill(hdr.cells, header, True)
        for c in hdr.cells:
            shade(c, "DCE6F1")
        row_flag(hdr, "tblHeader")
        row_flag(hdr, "cantSplit")

        for r in rows:
            row = t.add_row()
            vals = (list(r) + [""] * ncol)[:ncol]
            fill(row.cells, vals, False)
            row_flag(row, "cantSplit")

        set_col_widths(t, wcm)

        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(8)
        sp.paragraph_format.line_spacing = 1.0
        set_run_font(sp.add_run(""), size=4)
        return t

    # -- figures -----------------------------------------------------------
    def figure(self, image, width_cm, caption):
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(os.path.join(self.media, image), width=Cm(width_cm))
        self._caption("Hình", caption, above=False)

    # -- generated lists ---------------------------------------------------
    def toc_field(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        field(
            p,
            ' TOC \\o "1-4" \\h \\z \\u ',
            "Nhấn Ctrl+A rồi F9 để cập nhật mục lục.",
            size=BODY_PT,
        )

    def entry_list(self, entries):
        """A list of tables/figures: label + text .... page (PAGEREF)."""
        for label, text, bname in entries:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.line_spacing = 1.3
            pf.space_after = Pt(4)
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(2.0)
            pf.tab_stops.add_tab_stop(Cm(TEXT_WIDTH_CM), 2, 2)  # RIGHT, DOTS
            short = re.sub(r"\s*\[[0-9,\-\s]+\]\s*\.?\s*$", "", text).rstrip(" .")
            add_run(p, label + " ", size=BODY_PT, bold=False)
            add_rich(p, short, size=BODY_PT)
            add_run(p, "\t", size=BODY_PT)
            field(p, " PAGEREF %s \\h " % bname, "—", size=BODY_PT)


# ------------------------------------------------------------------- parsing
def parse_and_build(b, src):
    lines = src.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1
        if not line.strip() or line.startswith("#"):
            continue
        tag, _, rest = line.partition(" ")
        rest = rest.strip()

        if tag in ("H1", "H2", "H3", "H4"):
            b.heading(int(tag[1]), rest)
        elif tag == "H1P":
            b.heading(1, rest, page_break=True)
        elif tag == "PB":
            b.doc.add_page_break()
        elif tag == "SECTION":
            s = b.doc.add_section(WD_SECTION.NEW_PAGE)
            s.page_width, s.page_height = Cm(21.0), Cm(29.7)
            s.top_margin, s.bottom_margin = Cm(3.0), Cm(3.0)
            s.left_margin, s.right_margin = Cm(3.5), Cm(2.0)
            s.header_distance, s.footer_distance = Cm(1.5), Cm(1.5)
        elif tag == "P":
            b.para(rest)
        elif tag == "PN":
            b.para(rest, indent=False)
        elif tag == "PC":
            b.para(rest, indent=False, align="c")
        elif tag == "PCB":
            b.para(rest, indent=False, align="c", bold=True)
        elif tag == "PI":
            b.para(rest, indent=False, align="c", italic=True)
        elif tag == "BUL":
            b.bullet(rest)
        elif tag == "REF":
            b.reference(rest)
        elif tag == "TOC":
            b.toc_field()
        elif tag == "LISTTABLES":
            b.entry_list(b.tables)
        elif tag == "LISTFIGURES":
            b.entry_list(b.figures)
        elif tag == "FIG":
            parts = [x.strip() for x in rest.split("|")]
            b.figure(parts[0], float(parts[1]), parts[2])
        elif tag in ("TABLE", "TABLEX"):
            caption = rest if tag == "TABLE" else None
            widths, aligns, fontsize, header, rows = None, None, None, None, []
            while i < len(lines):
                l2 = lines[i].rstrip()
                i += 1
                if l2.startswith("ENDTABLE"):
                    break
                if not l2.strip():
                    continue
                t2, _, r2 = l2.partition(" ")
                r2 = r2.strip()
                if t2 == "COLW":
                    widths = [float(x) for x in r2.split(",")]
                elif t2 == "ALIGN":
                    aligns = [x.strip() for x in r2.split(",")]
                elif t2 == "FSIZE":
                    fontsize = float(r2)
                elif t2 == "H":
                    header = [c.strip() for c in r2.split("|")]
                elif t2 == "R":
                    rows.append([c.strip() for c in r2.split("|")])
            b.table(caption, header, rows, widths, aligns, fontsize)
        else:
            b.para(line.strip())
    return b


def enforce_font_everywhere(doc):
    """Give every run an explicit Times New Roman rFonts, incl. headers/footers."""
    bodies = [doc.element.body]
    # only parts that already exist -- touching sec.first_page_header etc. would
    # materialise unused header/footer parts in the package
    for part in doc.part.package.parts:
        try:
            if part.content_type.endswith('header+xml') or part.content_type.endswith('footer+xml'):
                bodies.append(part.element)
        except Exception:
            pass
    for b in bodies:
        for r in b.iter(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                r.insert(0, rpr)
            _rfonts(rpr)


def retheme(path):
    """Rewrite the theme's major/minor latin typefaces to Times New Roman."""
    import shutil
    import zipfile

    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/theme/theme1.xml":
                txt = data.decode("utf-8")
                txt = re.sub(r'(<a:(?:latin|ea|cs) typeface=")[^"]*(")', r"\1Times New Roman\2", txt)
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, path)


def build(src_path, out_path, media="media2"):
    doc = new_document()
    b = Builder(doc, media)
    src = open(src_path, encoding="utf-8").read()

    # Two passes: the first collects captions so the lists can be emitted,
    # the second builds the real document with the lists filled in.
    probe_doc = new_document()
    probe = Builder(probe_doc, media)
    parse_and_build(probe, src.replace("\nLISTTABLES", "\n#").replace("\nLISTFIGURES", "\n#"))
    b.tables_preview = probe.tables
    b.figures_preview = probe.figures

    orig_entry_list = b.entry_list
    state = {"n": 0}

    def entry_list_patched(_entries):
        state["n"] += 1
        orig_entry_list(b.tables_preview if state["n"] == 1 else b.figures_preview)

    b.entry_list = entry_list_patched
    parse_and_build(b, src)
    b.entry_list = orig_entry_list

    # covers unnumbered, front matter in roman numerals, body restarting at 1
    n = len(doc.sections)
    if n >= 3:
        section_page_numbers(doc.sections[0], "lowerRoman", start=1, show=False)
        section_page_numbers(doc.sections[1], "lowerRoman", start=1, show=True)
        for s in doc.sections[2:]:
            section_page_numbers(s, "decimal", start=1, show=True)
    elif n == 2:
        section_page_numbers(doc.sections[0], "lowerRoman", start=1, show=True)
        section_page_numbers(doc.sections[1], "decimal", start=1, show=True)
    else:
        section_page_numbers(doc.sections[0], "decimal", start=1, show=True)

    enforce_font_everywhere(doc)
    doc.save(out_path)
    retheme(out_path)
    return probe


if __name__ == "__main__":
    src_path, out_path = sys.argv[1], sys.argv[2]
    media = sys.argv[3] if len(sys.argv) > 3 else "media2"
    p = build(src_path, out_path, media)
    print("tables:", len(p.tables), "figures:", len(p.figures))
